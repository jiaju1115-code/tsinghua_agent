from __future__ import annotations

import shutil
from pathlib import Path
from typing import Any

from .common import dedupe_path, ensure_output_path
from .models import FilePlan


def _set_east_asia_font(run: Any, name: str = "Microsoft YaHei") -> None:
    from docx.oxml.ns import qn

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _shade_cell(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[min(index, len(widths_dxa) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _repeat_header_row(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _configure_styles(document: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("清问·TsingAsk V2 本地候选版")
    _set_east_asia_font(footer_run)
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(110, 110, 110)


def _add_title_block(document: Any, plan: FilePlan) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run(plan.metadata.get("template_name", "校园事务文件"))
    _set_east_asia_font(run)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("7F3C8D")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(plan.title)
    _set_east_asia_font(run)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    if plan.subtitle:
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(18)
        run = subtitle.add_run(plan.subtitle)
        _set_east_asia_font(run)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(90, 90, 90)

    meta_rows = []
    if plan.author:
        meta_rows.append(("作者", plan.author))
    for key in ("course", "department", "date"):
        if plan.metadata.get(key):
            meta_rows.append(({"course": "课程", "department": "院系", "date": "日期"}[key], str(plan.metadata[key])))
    if meta_rows:
        table = document.add_table(rows=len(meta_rows), cols=2)
        table.style = "Table Grid"
        _set_table_geometry(table, [1800, 7560])
        for row, (label, value) in zip(table.rows, meta_rows):
            row.cells[0].text = label
            row.cells[1].text = value
            _shade_cell(row.cells[0], "F2F4F7")
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for item in paragraph.runs:
                        _set_east_asia_font(item)


def create_docx(plan: FilePlan, output_path: str | Path | None = None, template_path: str | Path | None = None) -> Path:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt

    path = dedupe_path(ensure_output_path(plan.title, "docx", output_path))
    document = Document(str(template_path)) if template_path else Document()
    placed_sections: set[int] = set()
    if not template_path:
        _configure_styles(document)
        _add_title_block(document, plan)
    else:
        replacements = {"{{title}}": plan.title, "{{subtitle}}": plan.subtitle, "{{author}}": plan.author}
        for key, value in plan.metadata.items():
            if isinstance(value, (str, int, float)):
                replacements[f"{{{{{key}}}}}"] = str(value)
        document_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )
        for index, section in enumerate(plan.sections, 1):
            content = "\n".join([*section.paragraphs, *section.bullets])
            placeholders = (f"{{{{section_{index}}}}}", f"{{{{{section.heading}}}}}")
            for placeholder in placeholders:
                replacements[placeholder] = content
                if placeholder in document_text:
                    placed_sections.add(index - 1)
        replace_docx_text(document, replacements)

    for section_index, section in enumerate(plan.sections):
        if section_index in placed_sections:
            continue
        if section.heading:
            document.add_heading(section.heading, level=1)
        for text in section.paragraphs:
            paragraph = document.add_paragraph(text)
            for run in paragraph.runs:
                _set_east_asia_font(run)
        for text in section.bullets:
            paragraph = document.add_paragraph(text, style="List Bullet")
            for run in paragraph.runs:
                _set_east_asia_font(run)
        if section.table:
            column_count = max(len(row) for row in section.table)
            table = document.add_table(rows=len(section.table), cols=column_count)
            table.style = "Table Grid"
            base = 9360 // column_count
            widths = [base] * column_count
            widths[-1] += 9360 - sum(widths)
            _set_table_geometry(table, widths)
            _repeat_header_row(table.rows[0])
            for row_index, values in enumerate(section.table):
                for col_index, value in enumerate(values):
                    cell = table.cell(row_index, col_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.text = str(value)
                    if row_index == 0:
                        _shade_cell(cell, "F2F4F7")
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(2)
                        for run in paragraph.runs:
                            _set_east_asia_font(run)
                            if row_index == 0:
                                run.bold = True

    if plan.sources:
        document.add_heading("资料来源", level=1)
        for source in plan.sources:
            text = f"{source.get('title', '来源')}：{source.get('url', '')}".rstrip("：")
            paragraph = document.add_paragraph(text, style="List Bullet")
            for run in paragraph.runs:
                _set_east_asia_font(run)
                run.font.size = Pt(9)
    document.save(path)
    return path


def _replace_paragraph(paragraph: Any, replacements: dict[str, str]) -> int:
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        updated = updated.replace(old, new)
    if updated == original:
        return 0
    if not paragraph.runs:
        paragraph.add_run(updated)
    else:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    return 1


def replace_docx_text(document: Any, replacements: dict[str, str]) -> int:
    count = 0
    containers = [document, *[section.header for section in document.sections], *[section.footer for section in document.sections]]
    for container in containers:
        for paragraph in container.paragraphs:
            count += _replace_paragraph(paragraph, replacements)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        count += _replace_paragraph(paragraph, replacements)
    return count


def modify_docx(
    input_path: str | Path,
    *,
    replacements: dict[str, str] | None = None,
    append_plan: FilePlan | None = None,
    output_path: str | Path | None = None,
) -> tuple[Path, int]:
    from docx import Document

    source = Path(input_path).resolve()
    path = dedupe_path(ensure_output_path(f"{source.stem}_modified", "docx", output_path))
    shutil.copy2(source, path)
    document = Document(path)
    changed = replace_docx_text(document, replacements or {})
    if append_plan:
        for section in append_plan.sections:
            document.add_heading(section.heading, level=1)
            for text in section.paragraphs:
                document.add_paragraph(text)
            for text in section.bullets:
                document.add_paragraph(text, style="List Bullet")
        changed += len(append_plan.sections)
    document.save(path)
    return path, changed


def read_docx(path: str | Path) -> dict[str, Any]:
    from docx import Document

    document = Document(str(path))
    paragraphs = [item.text for item in document.paragraphs if item.text.strip()]
    headers = [paragraph.text for section in document.sections for paragraph in section.header.paragraphs if paragraph.text.strip()]
    footers = [paragraph.text for section in document.sections for paragraph in section.footer.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables:
        tables.append([[cell.text for cell in row.cells] for row in table.rows])
    return {
        "format": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "sections": len(document.sections),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables), "headers": headers, "footers": footers,
    }
